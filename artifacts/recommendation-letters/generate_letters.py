from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
OUTPUT_DIR = Path(__file__).resolve().parent
ASSET_DIR = ROOT / "public" / "hackathon"

FONT_NAME = "TH SarabunPSK"
NAVY = RGBColor(11, 37, 69)
MUTED = RGBColor(91, 105, 120)
BLACK = RGBColor(20, 24, 28)

STUDENTS = (
    "นางสาวกัญญาณัฐ หูชัยภูมิ",
    "นางสาวภัทรมน กุศลาไสยานนท์",
    "นายอุ่นตะวัน เสนพันธ์",
)


def set_run_font(
    run,
    *,
    size: float,
    color: RGBColor = BLACK,
    bold: bool = False,
    italic: bool = False,
) -> None:
    run.font.name = FONT_NAME
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT_NAME)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT_NAME)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT_NAME)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic


def shade_run(run, fill: str = "EAF2F8") -> None:
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    run._element.get_or_add_rPr().append(shd)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.35)
    section.bottom_margin = Cm(1.25)
    section.left_margin = Cm(2.25)
    section.right_margin = Cm(2.25)
    section.header_distance = Cm(0.7)
    section.footer_distance = Cm(0.7)

    normal = doc.styles["Normal"]
    normal.font.name = FONT_NAME
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT_NAME)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_NAME)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    normal.font.size = Pt(15.5)
    normal.font.color.rgb = BLACK
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(3)
    normal.paragraph_format.line_spacing = 1.08

    section.footer.is_linked_to_previous = False
    footer_paragraph = section.footer.paragraphs[0]
    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_paragraph.paragraph_format.space_after = Pt(0)
    footer_run = footer_paragraph.add_run(
        "THE NEXT DECADE HACKATHON 2026  •  PREVENTIVE & PREDICTIVE HEALTHCARE"
    )
    set_run_font(footer_run, size=8.5, color=MUTED)


def add_letterhead(doc: Document) -> None:
    event_logo = doc.add_paragraph()
    event_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    event_logo.paragraph_format.space_after = Pt(1)
    event_logo.add_run().add_picture(str(ASSET_DIR / "HackLogo.png"), width=Cm(5.0))

    event_name = doc.add_paragraph()
    event_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    event_name.paragraph_format.space_after = Pt(1)
    run = event_name.add_run("THE NEXT DECADE HACKATHON 2026")
    set_run_font(run, size=10, color=NAVY, bold=True)

    organizer_logos = doc.add_paragraph()
    organizer_logos.alignment = WD_ALIGN_PARAGRAPH.CENTER
    organizer_logos.paragraph_format.space_after = Pt(5)
    for logo_name in ("PS.png", "StemLike.png", "AMSA.png"):
        organizer_logos.add_run().add_picture(str(ASSET_DIR / logo_name), height=Cm(0.55))
        organizer_logos.add_run("   ")


def add_date(doc: Document) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run("3 สิงหาคม พ.ศ. 2569")
    set_run_font(run, size=15.5)


def add_title(doc: Document) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(7)
    run = paragraph.add_run("หนังสือรับรองและจดหมายแนะนำ")
    set_run_font(run, size=20, color=NAVY, bold=True)


def add_label_line(doc: Document, label: str, value: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(2)
    label_run = paragraph.add_run(f"{label}  ")
    set_run_font(label_run, size=15.5, bold=True)
    value_run = paragraph.add_run(value)
    set_run_font(value_run, size=15.5)


def add_body_paragraph(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.first_line_indent = Cm(1.25)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.08
    run = paragraph.add_run(text)
    set_run_font(run, size=15.5)


def add_signature(doc: Document) -> None:
    closing = doc.add_paragraph()
    closing.alignment = WD_ALIGN_PARAGRAPH.CENTER
    closing.paragraph_format.left_indent = Cm(8.0)
    closing.paragraph_format.space_before = Pt(2)
    closing.paragraph_format.space_after = Pt(18)
    set_run_font(closing.add_run("ขอแสดงความนับถือ"), size=15.5)

    name = doc.add_paragraph()
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name.paragraph_format.left_indent = Cm(8.0)
    name.paragraph_format.space_after = Pt(0)
    run = name.add_run("(บุณยสิทธิ์ ฟาง)")
    set_run_font(run, size=15.5, bold=True)

    role = doc.add_paragraph()
    role.alignment = WD_ALIGN_PARAGRAPH.CENTER
    role.paragraph_format.left_indent = Cm(8.0)
    role.paragraph_format.space_after = Pt(0)
    run = role.add_run("ผู้ก่อตั้ง PassionSeed")
    set_run_font(run, size=14.5)

    event = doc.add_paragraph()
    event.alignment = WD_ALIGN_PARAGRAPH.CENTER
    event.paragraph_format.left_indent = Cm(8.0)
    event.paragraph_format.space_after = Pt(0)
    set_run_font(
        event.add_run("ผู้จัดการแข่งขัน The Next Decade Hackathon 2026"),
        size=13.5,
        color=NAVY,
    )


def build_letter(student_name: str) -> Document:
    doc = Document()
    configure_document(doc)
    add_letterhead(doc)
    add_date(doc)
    add_title(doc)
    add_label_line(doc, "เรื่อง", f"รับรองและแนะนำ {student_name}")
    add_label_line(doc, "เรียน", "ผู้เกี่ยวข้อง")

    paragraphs = (
        f"ข้าพเจ้า บุณยสิทธิ์ ฟาง ผู้ก่อตั้ง PassionSeed และผู้จัดการแข่งขัน The Next Decade Hackathon 2026 ขอรับรองว่า {student_name} เป็นสมาชิกทีม “ผัดเห็ดใส่บลอกคอลี่” ซึ่งได้รับรางวัลรองชนะเลิศอันดับ 2 ระดับมัธยมศึกษา จากการแข่งขันภายใต้แนวคิด Preventive & Predictive Healthcare",
        "การแข่งขันครั้งนี้มีผู้เข้าร่วมมากกว่า 800 คน โดยทีมผัดเห็ดใส่บลอกคอลี่ได้รับคัดเลือกเป็นหนึ่งในห้าทีมระดับมัธยมศึกษาที่เข้าสู่รอบชิงชนะเลิศ และนำเสนอผลงานต่อคณะกรรมการเมื่อวันที่ 20 มิถุนายน พ.ศ. 2569 หลังผ่านกระบวนการศึกษาปัญหา พัฒนาและทดสอบต้นแบบ ตลอดจนปรับปรุงผลงานจากข้อมูลและข้อเสนอแนะ",
        f"การมีส่วนร่วมของ {student_name} และความสำเร็จของทีมสะท้อนถึงความสามารถในการทำงานร่วมกับผู้อื่น การคิดแก้ปัญหาอย่างเป็นระบบและสร้างสรรค์ ความรับผิดชอบ และความมุ่งมั่นในการพัฒนาผลงานภายใต้ข้อจำกัดของเวลา",
        f"ข้าพเจ้าจึงมีความยินดีแนะนำ {student_name} เพื่อใช้ประกอบแฟ้มสะสมผลงาน (Portfolio) และการสมัครเข้าศึกษาต่อ และเชื่อว่าทักษะและประสบการณ์ดังกล่าวจะเป็นพื้นฐานสำคัญต่อการเรียนรู้และการสร้างสรรค์ผลงานที่เป็นประโยชน์ในอนาคต",
    )

    for paragraph in paragraphs:
        add_body_paragraph(doc, paragraph)

    add_signature(doc)
    return doc


def file_slug(student_name: str) -> str:
    return student_name.replace("นางสาว", "").replace("นาย", "").strip().replace(" ", "_")


def main() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    for student_name in STUDENTS:
        output_path = OUTPUT_DIR / f"Letter_of_Recommendation_{file_slug(student_name)}.docx"
        build_letter(student_name).save(output_path)
        print(output_path)


if __name__ == "__main__":
    main()
